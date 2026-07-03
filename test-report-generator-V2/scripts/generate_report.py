#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试报告生成器 V2
生成简洁的测试报告Word文档,适合直接复制到邮件中
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class TestReportV2:
    """测试报告生成器 V2 - 简洁邮件风格"""
    
    def __init__(self):
        """初始化报告生成器"""
        self.doc = Document()
        self._setup_styles()
        self.detail_section_added = False
        self.detail_subsection_counter = 0
    
    def _setup_styles(self):
        """设置默认样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)
        
        # 设置段落间距
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
    
    def _add_heading(self, text, level=1):
        """添加标题
        
        Args:
            text: 标题文本
            level: 1=一级标题, 2=二级标题
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        
        return p
    
    def _add_paragraph(self, text, bold=False, color=None, font_size=10.5):
        """添加段落
        
        Args:
            text: 段落文本
            bold: 是否加粗
            color: 字体颜色 (RGB tuple)
            font_size: 字体大小
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(font_size)
        run.font.bold = bold
        
        if color:
            run.font.color.rgb = RGBColor(*color)
        
        return p
    
    def _add_link(self, text, url):
        """添加带链接的文本 (显示为蓝色)
        
        Args:
            text: 显示文本
            url: 链接地址
        """
        p = self.doc.add_paragraph()
        run = p.add_run(f"{text}: ")
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
        
        # 链接文本用蓝色
        link_run = p.add_run(url)
        link_run.font.name = '微软雅黑'
        link_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        link_run.font.size = Pt(10.5)
        link_run.font.color.rgb = RGBColor(0, 0, 255)
        link_run.font.underline = True
        
        return p
    
    def _add_inline_link(self, paragraph, text, url):
        """在段落中添加内联链接
        
        Args:
            paragraph: 段落对象
            text: 显示文本
            url: 链接地址
        """
        if text:
            run = paragraph.add_run(text)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10.5)
        
        if url:
            link_run = paragraph.add_run(url)
            link_run.font.name = '微软雅黑'
            link_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            link_run.font.size = Pt(10.5)
            link_run.font.color.rgb = RGBColor(0, 0, 255)
            link_run.font.underline = True
    
    def _add_simple_table(self, headers, data):
        """添加简单表格
        
        Args:
            headers: 表头列表
            data: 数据列表
        """
        if not headers or not data:
            return None
        
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            # 设置表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(9)
            # 设置背景色为浅灰
            self._set_cell_shading(cell, "E7E6E6")
        
        # 添加数据行
        for row_data in data:
            row = table.add_row()
            for i, value in enumerate(row_data):
                cell = row.cells[i]
                cell.text = str(value) if value is not None else ""
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        run.font.size = Pt(9)
        
        # 添加表格后的空行
        self.doc.add_paragraph()
        
        return table
    
    def _set_cell_shading(self, cell, color):
        """设置单元格背景色"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_image(self, image_path, width=5):
        """添加图片
        
        Args:
            image_path: 图片路径
            width: 图片宽度(英寸)
        """
        if not os.path.exists(image_path):
            self._add_paragraph(f"[图片未找到: {image_path}]", color=(255, 0, 0))
            return
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))
    
    def add_preface(self, requirement_desc, requirement_link):
        """添加前序: 需求描述和链接 (必须)
        
        Args:
            requirement_desc: 需求描述
            requirement_link: 需求链接
        """
        # 需求描述
        p = self.doc.add_paragraph()
        run1 = p.add_run("需求描述: ")
        run1.font.name = '微软雅黑'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(10.5)
        run1.font.bold = True
        
        run2 = p.add_run(requirement_desc)
        run2.font.name = '微软雅黑'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10.5)
        
        # 需求链接 (蓝色)
        self._add_link("需求链接", requirement_link)
        
        # 空行
        self.doc.add_paragraph()
    
    def add_conclusion(self, items=None, bug_table=None, remaining_issues=None, summary=None, problems=None):
        """添加测试结论 (必须)
        
        新格式 (推荐):
        Args:
            items: 测试项目列表，每项包含:
                   - type: 测试类型 (如"功能测试"、"效果测试"等)
                   - desc: 测试描述
                   - link: 可选链接
                   - sub_items: 可选子项目列表
            bug_table: 问题单列表 [{"id": "BUG001", "desc": "描述", "status": "已修复", "link": "xxx"}]
            remaining_issues: 遗留问题列表 [{"desc": "描述", "link": "xxx"}]
        
        旧格式 (兼容):
        Args:
            summary: 一句话总结
            problems: 问题列表 [{"id": "1", "desc": "问题", "status": "已解决", "reason": "备注"}]
        """
        self._add_heading("一、测试结论", level=1)
        
        # 新格式: 使用items列表
        if items:
            for i, item in enumerate(items, 1):
                item_type = item.get("type", "")
                item_desc = item.get("desc", "")
                item_link = item.get("link")
                sub_items = item.get("sub_items", [])
                
                # 创建段落: "1、功能测试：描述文字，链接"
                p = self.doc.add_paragraph()
                
                # 序号和类型
                run = p.add_run(f"{i}、{item_type}：")
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10.5)
                
                # 描述文字
                if item_desc:
                    desc_text = item_desc
                    if item_link:
                        desc_text += "，"
                    run2 = p.add_run(desc_text)
                    run2.font.name = '微软雅黑'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run2.font.size = Pt(10.5)
                
                # 链接 (蓝色)
                if item_link:
                    link_run = p.add_run(item_link)
                    link_run.font.name = '微软雅黑'
                    link_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    link_run.font.size = Pt(10.5)
                    link_run.font.color.rgb = RGBColor(0, 0, 255)
                    link_run.font.underline = True
                
                # 子项目 (如有)
                if sub_items:
                    for j, sub_item in enumerate(sub_items, 1):
                        sub_p = self.doc.add_paragraph()
                        sub_p.paragraph_format.left_indent = Pt(20)
                        sub_run = sub_p.add_run(f"   {j}) {sub_item}")
                        sub_run.font.name = '微软雅黑'
                        sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        sub_run.font.size = Pt(10.5)
        
        # 旧格式兼容: 使用summary
        elif summary:
            self._add_paragraph(summary)
        
        # 问题单表格 (新格式)
        if bug_table and len(bug_table) > 0:
            self.doc.add_paragraph()  # 空行
            self._add_paragraph("问题单修复情况：", bold=True)
            
            headers = ["问题单", "问题描述", "状态"]
            data = []
            for bug in bug_table:
                bug_id = bug.get("id", "")
                bug_link = bug.get("link", "")
                # 如果有link，显示link，否则显示id
                display_id = bug_link if bug_link else bug_id
                data.append([
                    display_id,
                    bug.get("desc", ""),
                    bug.get("status", "")
                ])
            
            self._add_simple_table(headers, data)
        
        # 遗留问题 (新格式)
        if remaining_issues and len(remaining_issues) > 0:
            p = self.doc.add_paragraph()
            run = p.add_run("遗留问题：")
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10.5)
            run.font.bold = True
            
            for i, issue in enumerate(remaining_issues, 1):
                issue_p = self.doc.add_paragraph()
                issue_desc = issue.get("desc", "")
                issue_link = issue.get("link")
                
                # 序号和描述
                run = issue_p.add_run(f"{i}.{issue_desc}")
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10.5)
                
                # 问题单链接
                if issue_link:
                    run2 = issue_p.add_run("，问题单：")
                    run2.font.name = '微软雅黑'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run2.font.size = Pt(10.5)
                    
                    link_run = issue_p.add_run(issue_link)
                    link_run.font.name = '微软雅黑'
                    link_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    link_run.font.size = Pt(10.5)
                    link_run.font.color.rgb = RGBColor(0, 0, 255)
                    link_run.font.underline = True
        
        # 旧格式问题列表表格 (兼容)
        if problems and len(problems) > 0:
            self._add_paragraph("问题列表:", bold=True)
            
            headers = ["序号", "问题描述", "状态", "备注"]
            data = []
            for prob in problems:
                data.append([
                    prob.get("id", ""),
                    prob.get("desc", ""),
                    prob.get("status", ""),
                    prob.get("reason", "-")
                ])
            
            self._add_simple_table(headers, data)
    
    def _ensure_detail_section(self):
        """确保详细数据章节已添加"""
        if not self.detail_section_added:
            self._add_heading("二、详细数据", level=1)
            self.detail_section_added = True
    
    def _get_next_subsection_number(self):
        """获取下一个子章节编号"""
        self.detail_subsection_counter += 1
        return f"2.{self.detail_subsection_counter}"
    
    def add_effect_data(self, description, table_headers=None, table_data=None, 
                        links=None, images=None):
        """添加效果测试数据
        
        Args:
            description: AI生成的数据描述
            table_headers: 表头列表
            table_data: 数据列表
            links: 链接列表
            images: 图片路径列表
        """
        self._ensure_detail_section()
        
        subsection = self._get_next_subsection_number()
        self._add_heading(f"{subsection} 效果测试数据", level=2)
        
        # 描述
        if description:
            self._add_paragraph(description)
        
        # 表格
        if table_headers and table_data:
            self._add_simple_table(table_headers, table_data)
        
        # 链接
        if links:
            for i, link in enumerate(links):
                self._add_link(f"数据链接{i+1}" if len(links) > 1 else "数据链接", link)
        
        # 图片
        if images:
            for img in images:
                self._add_image(img)
    
    def add_performance_data(self, description, table_headers=None, table_data=None,
                             links=None, images=None):
        """添加性能测试数据
        
        Args:
            description: AI生成的数据描述
            table_headers: 表头列表
            table_data: 数据列表
            links: 链接列表
            images: 图片路径列表
        """
        self._ensure_detail_section()
        
        subsection = self._get_next_subsection_number()
        self._add_heading(f"{subsection} 性能测试数据", level=2)
        
        # 描述
        if description:
            self._add_paragraph(description)
        
        # 表格
        if table_headers and table_data:
            self._add_simple_table(table_headers, table_data)
        
        # 链接
        if links:
            for i, link in enumerate(links):
                self._add_link(f"数据链接{i+1}" if len(links) > 1 else "数据链接", link)
        
        # 图片
        if images:
            for img in images:
                self._add_image(img)
    
    def add_stability_data(self, description, table_headers=None, table_data=None,
                           links=None, images=None):
        """添加稳定性测试数据
        
        Args:
            description: AI生成的数据描述
            table_headers: 表头列表
            table_data: 数据列表
            links: 链接列表
            images: 图片路径列表
        """
        self._ensure_detail_section()
        
        subsection = self._get_next_subsection_number()
        self._add_heading(f"{subsection} 稳定性测试数据", level=2)
        
        # 描述
        if description:
            self._add_paragraph(description)
        
        # 表格
        if table_headers and table_data:
            self._add_simple_table(table_headers, table_data)
        
        # 链接
        if links:
            for i, link in enumerate(links):
                self._add_link(f"数据链接{i+1}" if len(links) > 1 else "数据链接", link)
        
        # 图片
        if images:
            for img in images:
                self._add_image(img)
    
    def add_bypass_data(self, description, results=None, table_headers=None, 
                        table_data=None, links=None, images=None):
        """添加旁路测试数据
        
        Args:
            description: AI生成的数据描述
            results: 简短结果描述
            table_headers: 表头列表
            table_data: 数据列表
            links: 链接列表
            images: 图片路径列表
        """
        self._ensure_detail_section()
        
        subsection = self._get_next_subsection_number()
        self._add_heading(f"{subsection} 旁路测试数据", level=2)
        
        # 描述
        if description:
            self._add_paragraph(description)
        
        # 结果
        if results:
            self._add_paragraph(results)
        
        # 表格
        if table_headers and table_data:
            self._add_simple_table(table_headers, table_data)
        
        # 链接
        if links:
            for i, link in enumerate(links):
                self._add_link(f"数据链接{i+1}" if len(links) > 1 else "数据链接", link)
        
        # 图片
        if images:
            for img in images:
                self._add_image(img)
    
    def add_custom_section(self, title, description=None, table_headers=None, 
                           table_data=None, links=None, images=None):
        """添加自定义数据章节
        
        Args:
            title: 章节标题
            description: 描述
            table_headers: 表头列表
            table_data: 数据列表
            links: 链接列表
            images: 图片路径列表
        """
        self._ensure_detail_section()
        
        subsection = self._get_next_subsection_number()
        self._add_heading(f"{subsection} {title}", level=2)
        
        if description:
            self._add_paragraph(description)
        
        if table_headers and table_data:
            self._add_simple_table(table_headers, table_data)
        
        if links:
            for i, link in enumerate(links):
                self._add_link(f"链接{i+1}" if len(links) > 1 else "链接", link)
        
        if images:
            for img in images:
                self._add_image(img)
    
    def save(self, output_path):
        """保存报告
        
        Args:
            output_path: 输出文件路径
        """
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        self.doc.save(output_path)
        print(f"测试报告已生成: {os.path.abspath(output_path)}")
        return os.path.abspath(output_path)


def main():
    """示例用法"""
    report = TestReportV2()
    
    # 前序: 需求描述和链接 (必须)
    report.add_preface(
        requirement_desc="本次测试验证V6.0.0版本的智能检测功能优化",
        requirement_link="https://tapd.woa.com/xxx/story/detail/12345"
    )
    
    # 一、测试结论 (新格式 - 列表形式)
    report.add_conclusion(
        items=[
            {"type": "功能测试", "desc": "测试通过，用例地址", "link": "https://zhiyan.woa.com/testx/xxx"},
            {"type": "效果测试", "desc": "对比上版本(V5.8.0)主要优化通过率，反面通过率提升较明显"},
            {"type": "性能测试", "desc": "相比上版本(V5.8.0)，原子能力性能均有提升(提升约10%)"},
            {"type": "稳定性测试", "desc": "执行15小时未见明显异常，但有轻微内存泄漏", 
             "sub_items": ["T4服务独占显存13.5G；最大内存占用4.4G"]}
        ],
        bug_table=[
            {"id": "BUG001", "desc": "某场景下响应时间偏长", "status": "已修复", "link": "https://tapd.woa.com/xxx/bug/1"},
            {"id": "BUG002", "desc": "边缘case处理不完善", "status": "已修复", "link": "https://tapd.woa.com/xxx/bug/2"}
        ],
        remaining_issues=[
            {"desc": "内存泄漏问题待后续修复", "link": "https://tapd.woa.com/xxx/bug/3"}
        ]
    )
    
    # 2.1 效果测试数据
    report.add_effect_data(
        description="对比V6.0.0与V5.8.0版本的检测效果,V6.0.0在各项指标上均有提升。",
        table_headers=["版本", "回数", "请求成功数", "成功率", "检测通过率"],
        table_data=[
            ["V6.0.0", "10", "100", "100%", "95%"],
            ["V5.8.0", "10", "100", "100%", "90%"]
        ],
        links=["https://docs.qq.com/sheet/xxx"]
    )
    
    # 2.2 性能测试数据
    report.add_performance_data(
        description="性能压测结果显示,QPS达到100,平均响应时间50ms,P99为80ms。",
        table_headers=["path", "params", "tps", "avg_cost", "p99", "diff"],
        table_data=[
            ["/api/v1/detect", "{}", "100", "50ms", "80ms", "+5%"],
            ["/api/v1/analyze", "{}", "80", "60ms", "100ms", "+3%"]
        ],
        links=["https://grafana.xxx.com/dashboard"]
    )
    
    # 2.3 稳定性测试数据
    report.add_stability_data(
        description="稳定性测试运行7天,服务稳定无异常。",
        table_headers=["指标", "数值"],
        table_data=[
            ["运行时长", "7天"],
            ["请求总数", "1000000"],
            ["成功率", "99.99%"],
            ["平均响应时间", "52ms"]
        ]
    )
    
    # 保存
    report.save("测试报告_V2_示例.docx")


if __name__ == "__main__":
    main()
